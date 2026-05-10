from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

print("=== PROJECT STATEMENT ===")
try:
    doc = Document('CityMind_Project_Statement (2).docx')
    for i, t in enumerate(doc.tables):
        print(f"--- TABLE {i} ---")
        for row in t.rows:
            print(" | ".join([c.text.replace("\n", " ") for c in row.cells]))
except Exception as e:
    print(f"Error reading project statement: {e}")

print("=== PHASE 1 REPORT ===")
try:
    doc2 = Document('i240838_i240839_i240726_Phase1-report (1).docx')
    for i, t in enumerate(doc2.tables):
        print(f"--- TABLE {i} ---")
        for row in t.rows:
            print(" | ".join([c.text.replace("\n", " ") for c in row.cells]))
except Exception as e:
    print(f"Error reading phase 1 report: {e}")
