from docx import Document
import sys

sys.stdout.reconfigure(encoding='utf-8')

print("=== PROJECT STATEMENT TEXT ===")
try:
    doc = Document('CityMind_Project_Statement (2).docx')
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            print(text)
except Exception as e:
    print(f"Error: {e}")

print("\n\n=== PHASE 1 REPORT TEXT ===")
try:
    doc2 = Document('i240838_i240839_i240726_Phase1-report (1).docx')
    for p in doc2.paragraphs:
        text = p.text.strip()
        if text:
            print(text)
except Exception as e:
    print(f"Error: {e}")
